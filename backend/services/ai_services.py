import io
import asyncio
import json
import logging
import re
import base64
import pypdfium2 as pdfium
from docx import Document
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from core.config import Config
from core.database import supabase
from services.fraud_service import fraud_detector

logger = logging.getLogger(__name__)

# --- PROMPTS ---
PRE_AUTH_SYSTEM_PROMPT = """You are ClaimRidge AI, an expert Medical Officer and Fraud Detection agent for a health insurance company.
Your job is to review a pre-authorisation request based on multiple clinical documents submitted by the provider.

## EXPECTED CLAIM DETAILS (Submitted via Portal)
- Patient Name: **{expected_patient_name}**
- Patient ID: **{expected_patient_id}**
- Provider Name: **{expected_provider_name}**

## PAYER POLICY RULES (RAG Context)
{policy_rules}

## ANTI-FRAUD SYSTEM FLAGS
{fraud_context}

## YOUR TASK & CRITICAL RULES
1. **IDENTITY VERIFICATION**: You MUST verify that the patient name in the clinical documents matches the "Expected Claim Details" above.
2. **FRAUD ANALYSIS**: If any Anti-Fraud Flags are present, analyze the clinical documentation to see if they correlate with clinical inconsistencies (e.g., suspicious visit frequency, mismatched diagnosis). Provide a clear reason in your rationale.
3. **CLINICAL EVIDENCE CHECK**: If the documents are blank forms or generic instructions, "deny" the request.
4. **POLICY ADJUDICATION**: Cross-reference the clinical findings against the Payer Policy Rules. If the rules are short or generic (e.g. "Standard medical necessity guidelines apply"), USE YOUR OWN MEDICAL KNOWLEDGE to approve it if standard conservative treatments were attempted.
5. **DECISION**: Render a decision: "approve", "deny", or "escalate".
   - **ESCALATE** if fraud is strongly suspected or documents explicitly contradict each other.

Respond ONLY with valid JSON in this exact format:
{{
    "decision": "approve" | "deny" | "escalate",
    "rationale": "Provide a clear justification, including fraud analysis if applicable."
}}
"""

# --- LLM INITIALIZATION ---
def get_vision_llm():
    return ChatGoogleGenerativeAI(
        model=Config.OCR_MODEL, 
        google_api_key=Config.GEMINI_API_KEY, 
        temperature=0.1
    )

def get_llm():
    return ChatGroq(
        api_key=Config.GROQ_API_KEY, 
        model_name=Config.LLM_MODEL, 
        temperature=0.1
    )

def get_embeddings():
    return GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001", 
        google_api_key=Config.GEMINI_API_KEY
    )

# --- UTILITIES ---
def extract_json_from_text(text: str) -> str:
    text = re.sub(r'```[\w]*\s*\n?', '', text)
    start_idx = text.find('{')
    end_idx = text.rfind('}')
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        return text[start_idx:end_idx + 1]
    return text.strip()

async def extract_text_from_file(base64_data: str, media_type: str) -> str:
    """Extracts text from PDF, Images, or Word documents with high-fidelity formatting."""
    
    # 1. Handle Word Documents (.docx)
    if "word" in media_type or "officedocument" in media_type:
        try:
            doc_bytes = base64.b64decode(base64_data)
            doc = Document(io.BytesIO(doc_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            return "Error: Could not read Word document."

    # 2. Handle PDFs and Images via Gemini Vision OCR
    vision_llm = get_vision_llm()
    
    # ENHANCED EXTRACTION PROMPT
    vision_prompt = """
    You are a world-class medical document transcriptionist. 
    Your goal is to extract clinical data from the provided image/document and format it for EXTREME READABILITY.

    ### FORMATTING RULES (STRICT):
    1. MAIN TITLES: Put category names in bold followed by a colon (e.g., **Patient Identification:**).
    2. DATA FIELDS: List sub-titles (labels) followed by their info on separate lines under the main title.
    3. LINE BREAKS: Use a double line break between categories and a single line break between every data field.
       Example:
       **Patient Identification:**
       Medicare Coverage: Yes
       Patient Name: XX YY.

       **Insured Information:**
       Insured's I.D. Number: 987 65 4321A
    
    4. CLEAN UP: IGNORE all boilerplate instructions, footer codes, and "Leave Blank" fields.
    5. NO BOILERPLATE: Do not include text like "DO NOT WRITE IN THIS SPACE" or generic form titles.

    If the document is a blank form with no handwritten or typed data, output exactly: "[NO CLINICAL DATA FOUND - BLANK FORM]"
    """
    
    messages = [
        HumanMessage(content=[
            {"type": "text", "text": vision_prompt},
            {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{base64_data}"}}
        ])
    ]
    response = await vision_llm.ainvoke(messages)
    return response.content

# --- DEDICATED FRAUD MODEL (NATIVE INTEGRATION) ---
async def check_fraud_system(request_data: dict) -> dict:
    """
    Calls the native XGBoost Fraud Service directly in memory.
    """
    try:
        # Native Python call. Executes in microseconds.
        return await fraud_detector.analyze_claim(request_data)
    except Exception as e:
        logger.error(f"Native Fraud Module failed: {e}. Bypassing statistical check.")
        return {
            "risk_level": "low",
            "fraud_score": 0.0,
            "flags": []
        }


async def process_pre_auth_case(pre_auth_id: str, insurer_id: str, attachments: list):
    """Background task to handle OCR extraction and AI evaluation."""
    logger.info(f"Starting background processing for Pre-Auth: {pre_auth_id}")
    
    # 1. Perform OCR on all documents
    for att in attachments:
        try:
            # att is a dict with file_name, content_type, content (base64)
            extracted_text = await extract_text_from_file(att["content"], att["content_type"])
            
            # Update the existing document record with extracted text
            supabase.table("pre_auth_documents").update({
                "extracted_text": extracted_text
            }).eq("pre_auth_id", pre_auth_id).eq("file_name", att["file_name"]).execute()
            
        except Exception as e:
            logger.error(f"Background OCR failed for {att['file_name']}: {e}")

    # 2. Run the actual clinical evaluation
    await evaluate_pre_auth(pre_auth_id, insurer_id)

# --- CORE SERVICES ---
async def evaluate_pre_auth(pre_auth_id: str, insurer_id: str):
    """Orchestrates the 2-Layer Defense: Fraud Model -> RAG Policy -> LLM Clinical Triage."""
    logger.info(f"Starting Evaluation Pipeline for Pre-Auth: {pre_auth_id}")
    
    # 1. Fetch the Request Details
    req_res = supabase.table("pre_auth_requests").select("*").eq("id", pre_auth_id).execute()
    if not req_res.data:
        logger.error(f"Pre-auth request {pre_auth_id} not found.")
        return
    request_data = req_res.data[0]

    fraud_result = await check_fraud_system(request_data)
    fraud_context = "No significant statistical fraud flags detected."
    
    if fraud_result["risk_level"] == "high":
        flags_str = ", ".join(fraud_result["flags"])
        fraud_context = f"CRITICAL WARNING: The anti-fraud system has flagged this claim as HIGH RISK (Score: {fraud_result['fraud_score']}%). Flags: {flags_str}. Please scrutinize the clinical documents for inconsistencies."
    elif fraud_result["risk_level"] == "medium":
        flags_str = ", ".join(fraud_result["flags"])
        fraud_context = f"NOTICE: Medium risk flags detected: {flags_str}."

    # ==========================================
    # LAYER 2: CLINICAL LLM TRIAGE
    # ==========================================
    # 2. Fetch all documents for this case
    docs_res = supabase.table("pre_auth_documents").select("file_name, extracted_text").eq("pre_auth_id", pre_auth_id).execute()
    
    if not docs_res.data:
        logger.error("No documents found for this pre-auth.")
        # If no docs, we can't evaluate, but we should at least update status to show we tried
        supabase.table("pre_auth_requests").update({
            "status": "escalated",
            "ai_decision": "escalate",
            "ai_rationale": "No documents found for clinical review."
        }).eq("id", pre_auth_id).execute()
        return

    # 3. Combine all documents into one context window
    combined_clinical_context = "=== CLINICAL DOCUMENTS ===\n\n"
    for doc in docs_res.data:
        combined_clinical_context += f"--- Document: {doc['file_name']} ---\n{doc['extracted_text']}\n\n"

    # 4. RAG: Fetch relevant policy rules
    embeddings_model = get_embeddings()
    query_vector = embeddings_model.embed_query(combined_clinical_context[:2000]) 
    
    policy_res = supabase.rpc("match_policy_rules", {
        "query_embedding": query_vector,
        "match_threshold": 0.3,
        "match_count": 5,
        "p_insurer_id": insurer_id
    }).execute()
    
    policy_rules = "Standard medical necessity guidelines apply."
    if policy_res.data:
        policy_rules = "\n".join([match["content"] for match in policy_res.data])

    # 5. Ask the LLM to reason (Injecting Expected Details to prevent Identity Fraud)
    llm = get_llm()
    prompt = PRE_AUTH_SYSTEM_PROMPT.format(
        expected_patient_name=request_data.get("patient_name", "Unknown"),
        expected_patient_id=request_data.get("patient_id", "Unknown"),
        expected_provider_name=request_data.get("provider_name", "Unknown"),
        policy_rules=policy_rules,
        fraud_context=fraud_context
    )
    
    messages = [
        SystemMessage(content=prompt),
        HumanMessage(content=combined_clinical_context)
    ]
    
    response = await llm.ainvoke(messages)
    json_str = extract_json_from_text(response.content)
    
    try:
        result = json.loads(json_str)
        decision = result.get("decision", "escalate")
        rationale = result.get("rationale", "Failed to parse rationale.")
        
        # 6. Update Database
        new_status = 'escalated' if decision == 'escalate' else decision
        
        supabase.table("pre_auth_requests").update({
            "status": new_status,
            "ai_decision": decision,
            "ai_rationale": rationale
        }).eq("id", pre_auth_id).execute()
        
        logger.info(f"Pre-Auth {pre_auth_id} evaluated. Decision: {decision}")
        
    except Exception as e:
        logger.error(f"Failed to parse AI Pre-Auth Decision: {e}")
        supabase.table("pre_auth_requests").update({
            "status": "escalated",
            "ai_decision": "escalate",
            "ai_rationale": f"System error during AI evaluation. Manual review required."
        }).eq("id", pre_auth_id).execute()

async def process_and_embed_policy_file(insurer_id: str, base64_data: str, file_name: str = ""):
    """Called once when an insurer uploads their policy. Chunks the PDF/Word and saves to Supabase."""
    logger.info(f"Starting policy file processing for insurer {insurer_id}")
    
    full_text = ""
    is_word = file_name.lower().endswith(".docx") or "officedocument" in base64_data[:50] 

    try:
        if is_word:
            # Handle Word
            doc_bytes = base64.b64decode(base64_data)
            doc = Document(io.BytesIO(doc_bytes))
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            # Handle PDF (Default)
            pdf_bytes = base64.b64decode(base64_data)
            pdf = pdfium.PdfDocument(pdf_bytes)
            for page in pdf:
                text_page = page.get_textpage()
                full_text += text_page.get_text_range() + "\n\n"
    except Exception as e:
        logger.error(f"Failed to parse policy file: {e}")
        raise ValueError(f"Could not read the uploaded file. Please ensure it is a valid PDF or .docx file. Error: {str(e)}")
        
    if not full_text.strip():
        logger.warning(f"Insurer {insurer_id} uploaded a file with no readable text.")
        raise ValueError("The uploaded file contains no readable text.")

    # 2. Chunk text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        length_function=len,
    )
    chunks = text_splitter.split_text(full_text)
    logger.info(f"Split PDF into {len(chunks)} chunks. Beginning embedding...")

    embeddings_model = get_embeddings()
    
    # 3. Clear old policy chunks for this insurer
    supabase.table("policy_chunks").delete().eq("insurer_id", insurer_id).execute()
    
    # 4. Embed and insert in safe batches to respect rate limits
    batch_size = 50 
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i+batch_size]
        
        try:
            vectors = embeddings_model.embed_documents(batch_chunks)
        except Exception as e:
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                logger.warning("Rate limit hit during embedding. Retrying in 10 seconds...")
                await asyncio.sleep(10) 
                vectors = embeddings_model.embed_documents(batch_chunks)
            else:
                raise e
        
        payload = []
        for chunk, vector in zip(batch_chunks, vectors):
            payload.append({
                "insurer_id": insurer_id,
                "content": chunk,
                "embedding": vector
            })
            
        supabase.table("policy_chunks").insert(payload).execute()
        logger.info(f"Inserted batch {i//batch_size + 1} into database.")
        
        await asyncio.sleep(3)
        
    logger.info(f"Finished processing policy for insurer {insurer_id}. Total {len(chunks)} chunks saved.")